import json
import re

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def create_qa_file(raw_qa_file, delimiter="\t", output_file="qa_for_tonic_validate.json"):
    """
    Creates a JSON file for Tonic Validate evaluation from a raw Q/A DOCX file,
    extracting reference article and reference text with improved heuristics.

    Args:
        raw_qa_file (str): Path to the raw Q/A DOCX file.
        delimiter (str, optional): Delimiter used in the raw file (default: "\t").
            This argument is not used in this implementation.
        output_file (str, optional): Path to the output JSON file (default: "qa_for_tonic_validate.json").
    """

    qa_data = []
    with open(raw_qa_file, 'rb') as f:
        document = Document(f)

        # Track current question to associate with following answer(s)
        current_question = None

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            # Identify question based on starting with "?" or ending with a colon (":")
            if text.endswith(":"):
                current_question = text.strip()
            elif text.startswith("?") or current_question:
                if current_question:
                    # Answer paragraph after a question
                    answer = text

                    # Extract reference information (heuristics-based, improved patterns)
                    reference_article = None
                    reference_text = None
                    reference_patterns = [
                        r"(see also|relevant source|reference): (.*?)$",  # End of paragraph
                        r"(article|further details): (.*?)\n",  # Followed by newline
                        r"\[([^\]]*)\]\(([^)]*)\)",  # Link in brackets (URL, text)
                    ]
                    for pattern in reference_patterns:
                        match = re.search(pattern, answer, re.DOTALL)
                        if match:
                            reference_article = match.group(1).strip() if match.group(1) else None
                            reference_text = match.group(2).strip() if match.group(2) else None
                            break  # Stop after finding a match

                    qa_data.append({
                        "question": current_question,
                        "answer": answer,
                        "reference_article": reference_article,
                        "reference_text": reference_text
                    })

            # Clear current_question if not followed by another question indicator
            if not (text.startswith("?") or text.endswith(":")):
                current_question = None

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(qa_data, f, indent=2)  # Indent for readability



# def extract_text_from_docx(docx_file):
#     document = Document(docx_file)
#     text = []
#     for paragraph in document.paragraphs:
#         text.append(paragraph.text)
#     return text

def extract_hyperlinks_from_docx(docx_file):
    document = Document(docx_file)
    hyperlinks = {}
    for i, rel in enumerate(document.part.rels.values()):
        if rel.reltype == RT.HYPERLINK:
            hyperlinks[i] = rel.target_ref
    return hyperlinks

def extract_qa(raw_qa, hyperlinks):
    qa_list = []
    current_qa = None

    for i, line in enumerate(raw_qa):
        if line.startswith("Q:") or line.endswith("?"):
            if current_qa is not None:
                # If there's a hyperlink associated with this Q/A pair, assign it to the "reference_article" field
                if i in hyperlinks:
                    current_qa["reference_article"] = hyperlinks[i]
                qa_list.append(current_qa)
            current_qa = {"index": i, "question": line.lstrip("Q:").strip()}
        elif line.startswith("A:"):
            answer = line[3:].strip()
            current_qa["answer"] = answer

    # Add the last Q/A pair to the list
    if current_qa is not None:
        # If there's a hyperlink associated with this Q/A pair, assign it to the "reference_article" field
        if current_qa["index"] in hyperlinks:
            current_qa["reference_article"] = hyperlinks[current_qa["index"]]
        qa_list.append(current_qa)

    return qa_list

def extract_text_from_docx(docx_file):
    document = Document(docx_file)
    qa_pairs = []
    current_qa = None

    for i, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("Q:") or text.endswith("?"):
            if current_qa is not None:
                qa_pairs.append(current_qa)
            current_qa = {"question": text.lstrip("Q:").strip(), "answer": ""}
        elif text.startswith("A:"):
            if current_qa is None:
                continue  # Skip answers without questions
            current_qa["answer"] = text.lstrip("A:").strip()

        # Check for hyperlinks in the paragraph
        if paragraph.hyperlinks:
            for hyperlink in paragraph.hyperlinks:
                if current_qa is not None:
                    current_qa["reference_article"] = hyperlink.address
                    break  # Break after finding the first hyperlink in the paragraph

    # Add the last Q/A pair to the list
    if current_qa is not None:
        qa_pairs.append(current_qa)

    return qa_pairs


def save_to_json(qa_list, output_file):
    with open(output_file, 'w', encoding='utf-8') as json_file:
        json.dump(qa_list, json_file, ensure_ascii=False, indent=4)


if __name__ == "__main__":
    raw_qa_file = "raw_qa_files/Sourcegraph_Interview_Process_FAQ.docx"

    # Read raw Q/A data from the .docx file
    qa_list = extract_text_from_docx(raw_qa_file)

    # Save to JSON file
    save_to_json(qa_list, 'qa_for_tonic_validate.json')

# create_qa_file(raw_qa_file, output_file="qa_for_tonic_validate.json")
